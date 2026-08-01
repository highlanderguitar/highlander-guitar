from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import sys
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
R = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def main() -> None:
    source = Path(sys.argv[1]).resolve()
    output = Path(sys.argv[2]).resolve()
    media_dir = output / "source_media"
    output.mkdir(parents=True, exist_ok=True)
    media_dir.mkdir(parents=True, exist_ok=True)

    document = Document(source)
    rels = document.part.rels
    page = 1
    rows: list[dict] = []
    image_number = 0
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = paragraph.text.strip()
        image_refs = []
        for blip in paragraph._p.iter(A + "blip"):
            rel_id = blip.get(R + "embed")
            if not rel_id or rel_id not in rels:
                continue
            image_number += 1
            part = rels[rel_id].target_part
            suffix = Path(str(part.partname)).suffix or ".bin"
            filename = f"image-{image_number:04d}{suffix}"
            (media_dir / filename).write_bytes(part.blob)
            image_refs.append(filename)
        timestamps = re.findall(r"(?<!\d)(?:\d{1,2}:)?\d{1,2}:\d{2}(?!\d)", text)
        if text or image_refs:
            rows.append({
                "paragraph": index,
                "estimated_page": page,
                "style": paragraph.style.name if paragraph.style else "",
                "text": text,
                "timestamps": " | ".join(timestamps),
                "images": " | ".join(image_refs),
            })
        rendered_breaks = sum(1 for node in paragraph._p.iter(W + "lastRenderedPageBreak"))
        explicit_breaks = sum(1 for node in paragraph._p.iter(W + "br") if node.get(W + "type") == "page")
        page += rendered_breaks + explicit_breaks

    with (output / "paragraph_inventory.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=rows[0].keys())
        writer.writeheader()
        writer.writerows(rows)

    tables = []
    for table_index, table in enumerate(document.tables, 1):
        tables.append({
            "table": table_index,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "text": [[cell.text for cell in row.cells] for row in table.rows],
        })

    sha = hashlib.sha256(source.read_bytes()).hexdigest()
    summary = {
        "source": str(source),
        "sha256": sha,
        "bytes": source.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "nonempty_or_image_paragraphs": len(rows),
        "estimated_pages_from_saved_breaks": page,
        "tables": tables,
        "embedded_image_occurrences": image_number,
        "unique_extracted_images": len(list(media_dir.iterdir())),
        "styles": Counter(row["style"] for row in rows),
    }
    (output / "structural_summary.json").write_text(json.dumps(summary, indent=2, default=dict), encoding="utf-8")
    print(json.dumps({key: value for key, value in summary.items() if key not in {"tables", "styles"}}, indent=2))


if __name__ == "__main__":
    main()
